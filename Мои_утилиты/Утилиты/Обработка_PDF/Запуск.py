# -*- coding: utf-8 -*-
"""
Обработка PDF: текст (нативный + OCR), сводка по страницам, экспорт Excel / Word / JSON / CSV.
"""

from __future__ import annotations

import csv
import json
import os
import sys
import threading
import time
import tkinter as tk
from dataclasses import asdict
from tkinter import filedialog, messagebox, ttk
from typing import Any

_утил_dir = os.path.dirname(os.path.abspath(__file__))
if _утил_dir not in sys.path:
    sys.path.insert(0, _утил_dir)

_ядро = os.path.normpath(os.path.join(_утил_dir, "..", "..", "Ядро"))
if _ядро not in sys.path:
    sys.path.insert(0, _ядро)

import hub_theme as _T  # type: ignore
import оформление_утилиты as _shell  # type: ignore

import pdf_обработка as pdf_mod

_OPENPYXL_OK = False
try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment
    from openpyxl.worksheet.table import Table, TableStyleInfo

    _OPENPYXL_OK = True
except ImportError:
    Workbook = None  # type: ignore[misc, assignment]

_DOCX_OK = False
try:
    from docx import Document

    _DOCX_OK = True
except ImportError:
    Document = None  # type: ignore[misc, assignment]

_SETTINGS_NAME = "настройки_обработка_pdf.json"
_SETTINGS_PATH = os.path.join(_утил_dir, _SETTINGS_NAME)


def _rep_ok(родитель: tk.Misc):
    return getattr(родитель, "report_tab_error", lambda *_: None)


def _load_settings() -> dict[str, Any]:
    if not os.path.isfile(_SETTINGS_PATH):
        return {}
    try:
        with open(_SETTINGS_PATH, "r", encoding="utf-8") as f:
            return dict(json.load(f))
    except Exception:
        return {}


def _save_settings(data: dict[str, Any]) -> None:
    try:
        path = _SETTINGS_PATH
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(tmp, path)
    except Exception:
        pass


def batch_to_json_obj(batch: pdf_mod.BatchOutcome) -> dict[str, Any]:
    return {
        "schemaVersion": 1,
        "generatedAt": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "options": getattr(batch, "options_snapshot", {}),
        "warnings": [asdict(w) for w in batch.warnings],
        "files": [
            {
                "path": f.path,
                "page_count": f.page_count,
                "md5": f.md5,
                "metadata": f.metadata,
                "file_error": f.file_error,
                "duration_sec": f.duration_sec,
                "pages": [asdict(p) for p in f.pages],
            }
            for f in batch.files
        ],
    }


def export_excel(path: str, batch: pdf_mod.BatchOutcome) -> tuple[bool, str]:
    if not _OPENPYXL_OK or Workbook is None:
        return False, "Установите openpyxl: pip install openpyxl"
    wb = Workbook()
    # --- Страницы ---
    ws1 = wb.active
    ws1.title = "Страницы"
    h1 = [
        "Файл",
        "Стр.",
        "Источник",
        "Символов",
        "Превью",
        "Аннотации",
        "Картинок",
        "Печать?",
        "OCR conf",
        "Таблиц",
        "Превью таблицы",
        "Ошибка",
    ]
    for c, name in enumerate(h1, start=1):
        ws1.cell(row=1, column=c, value=name)
    row = 2
    for fo in batch.files:
        for pg in fo.pages:
            ws1.cell(row=row, column=1, value=os.path.basename(fo.path))
            ws1.cell(row=row, column=2, value=pg.page_index + 1)
            ws1.cell(row=row, column=3, value=pg.source)
            ws1.cell(row=row, column=4, value=pg.char_count)
            ws1.cell(row=row, column=5, value=pg.preview)
            ws1.cell(row=row, column=6, value=pg.annotation_types)
            ws1.cell(row=row, column=7, value=pg.image_count)
            ws1.cell(row=row, column=8, value="да" if pg.stamp_heuristic else "нет")
            ws1.cell(
                row=row,
                column=9,
                value=round(pg.ocr_confidence, 1) if pg.ocr_confidence is not None else "",
            )
            ws1.cell(row=row, column=10, value=pg.tables_found)
            ws1.cell(row=row, column=11, value=pg.table_preview)
            ws1.cell(row=row, column=12, value=pg.error or "")
            row += 1
    try:
        tab1 = Table(displayName="tbl_pages", ref=f"A1:L{max(1, row - 1)}")
        tab1.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        ws1.add_table(tab1)
    except Exception:
        pass
    for col in ws1.columns:
        for cell in col[1:]:
            try:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
            except Exception:
                pass
    # --- Файлы ---
    ws2 = wb.create_sheet("Файлы")
    h2 = ["Файл", "Страниц в PDF", "MD5", "Title", "Author", "Subject", "Ошибка файла", "Сек"]
    for c, name in enumerate(h2, start=1):
        ws2.cell(row=1, column=c, value=name)
    r = 2
    for fo in batch.files:
        meta = fo.metadata or {}
        ws2.cell(row=r, column=1, value=fo.path)
        ws2.cell(row=r, column=2, value=fo.page_count)
        ws2.cell(row=r, column=3, value=fo.md5 or "")
        ws2.cell(row=r, column=4, value=str(meta.get("title") or ""))
        ws2.cell(row=r, column=5, value=str(meta.get("author") or ""))
        ws2.cell(row=r, column=6, value=str(meta.get("subject") or ""))
        ws2.cell(row=r, column=7, value=fo.file_error or "")
        ws2.cell(row=r, column=8, value=round(fo.duration_sec, 2))
        r += 1
    # --- Предупреждения ---
    ws3 = wb.create_sheet("Предупреждения")
    for c, name in enumerate(["Тип", "Файл", "Страница", "Деталь"], start=1):
        ws3.cell(row=1, column=c, value=name)
    rr = 2
    for w in batch.warnings:
        ws3.cell(rr, 1, w.kind)
        ws3.cell(rr, 2, w.path)
        ws3.cell(rr, 3, w.page or "")
        ws3.cell(rr, 4, w.detail)
        rr += 1
    try:
        wb.save(path)
        return True, ""
    except Exception as e:
        return False, str(e)


def export_csv(path: str, batch: pdf_mod.BatchOutcome) -> tuple[bool, str]:
    try:
        with open(path, "w", encoding="utf-8-sig", newline="") as f:
            w = csv.writer(f, delimiter=";")
            w.writerow(
                [
                    "Файл",
                    "Стр.",
                    "Источник",
                    "Символов",
                    "Превью",
                    "Ошибка",
                ]
            )
            for fo in batch.files:
                for pg in fo.pages:
                    w.writerow(
                        [
                            fo.path,
                            pg.page_index + 1,
                            pg.source,
                            pg.char_count,
                            pg.preview,
                            pg.error or "",
                        ]
                    )
        return True, ""
    except Exception as e:
        return False, str(e)


def export_word(path: str, batch: pdf_mod.BatchOutcome, short: bool) -> tuple[bool, str]:
    if not _DOCX_OK or Document is None:
        return False, "Установите python-docx: pip install python-docx"
    try:
        doc = Document()
        doc.add_heading("Обработка PDF", 0)
        doc.add_paragraph(
            f"Сформировано: {time.strftime('%Y-%m-%d %H:%M')}. "
            f"Файлов: {len(batch.files)}."
        )
        for fo in batch.files:
            doc.add_heading(os.path.basename(fo.path), level=1)
            if fo.file_error:
                doc.add_paragraph(f"Ошибка: {fo.file_error}")
                continue
            for pg in fo.pages:
                doc.add_heading(f"Страница {pg.page_index + 1}", level=2)
                if short:
                    doc.add_paragraph(pg.preview or "(пусто)")
                else:
                    txt = pg.text or ""
                    if len(txt) > 120_000:
                        doc.add_paragraph(txt[:120_000] + "\n… [обрезано; полный текст в Excel/JSON]")
                    else:
                        for line in txt.split("\n"):
                            doc.add_paragraph(line)
        doc.save(path)
        return True, ""
    except Exception as e:
        return False, str(e)


def export_json(path: str, batch: pdf_mod.BatchOutcome) -> tuple[bool, str]:
    try:
        payload = batch_to_json_obj(batch)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        return True, ""
    except Exception as e:
        return False, str(e)


def запустить(родитель: tk.Misc) -> None:
    set_err = _rep_ok(родитель)

    def clear_tab_err() -> None:
        try:
            set_err(False)
        except Exception:
            pass

    _, root = _shell.карточка_утилиты(родитель, "Обработка PDF")

    settings = _load_settings()
    files_list: list[str] = []

    batch_holder: list[pdf_mod.BatchOutcome | None] = [None]
    cancel_ev = threading.Event()
    worker: list[threading.Thread | None] = [None]

    top_fr = tk.Frame(root, bg=_T.CARD)
    top_fr.pack(fill=tk.X, pady=(0, 6))

    def add_files() -> None:
        init_dir = settings.get("last_dir") or os.path.expanduser("~")
        paths = filedialog.askopenfilenames(
            title="PDF файлы",
            filetypes=[("PDF", "*.pdf"), ("Все", "*.*")],
            initialdir=init_dir if os.path.isdir(init_dir) else None,
        )
        if not paths:
            return
        settings["last_dir"] = os.path.dirname(os.path.abspath(paths[0]))
        for p in paths:
            ap = os.path.abspath(p)
            if ap.lower().endswith(".pdf") and ap not in files_list:
                files_list.append(ap)
        refresh_file_list()

    def add_folder() -> None:
        init_dir = settings.get("last_dir") or os.path.expanduser("~")
        d = filedialog.askdirectory(title="Папка с PDF", initialdir=init_dir if os.path.isdir(init_dir) else None)
        if not d:
            return
        settings["last_dir"] = d
        for name in sorted(os.listdir(d)):
            if name.lower().endswith(".pdf"):
                ap = os.path.join(d, name)
                if ap not in files_list:
                    files_list.append(ap)
        refresh_file_list()

    ttk.Button(top_fr, text="Файлы…", command=add_files).pack(side=tk.LEFT, padx=(0, 6))
    ttk.Button(top_fr, text="Папка…", command=add_folder).pack(side=tk.LEFT, padx=(0, 6))
    ttk.Button(
        top_fr,
        text="Очистить",
        command=lambda: (files_list.clear(), refresh_file_list()),
        style="Secondary.TButton",
    ).pack(side=tk.LEFT, padx=(0, 6))

    list_fr = tk.Frame(root, bg=_T.CARD)
    list_fr.pack(fill=tk.BOTH, expand=True, pady=(0, 8))
    sc = ttk.Scrollbar(list_fr)
    sc.pack(side=tk.RIGHT, fill=tk.Y)
    lb = tk.Listbox(
        list_fr,
        height=6,
        font=_T.FONT_MONO,
        yscrollcommand=sc.set,
        bg=_T.CARD_ALT,
        fg=_T.TEXT,
        selectbackground=_T.ACCENT,
        selectforeground=_T.ACCENT_FG,
        relief=tk.FLAT,
        highlightthickness=1,
        highlightbackground=_T.BORDER,
    )
    lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    sc.config(command=lb.yview)

    def refresh_file_list() -> None:
        lb.delete(0, tk.END)
        pw = pwd_var.get().strip() or None
        for p in files_list:
            bn = os.path.basename(p)
            n = ""
            if pdf_mod.FITZ_AVAILABLE:
                try:
                    d, _e = pdf_mod.open_pdf_authenticated(p, pw)
                    if d is not None:
                        n = str(d.page_count)
                        d.close()
                except Exception:
                    n = "?"
            lb.insert(tk.END, f"{bn}  ({n} стр.)")

    cfg_fr = ttk.LabelFrame(root, text="Режим")
    cfg_fr.pack(fill=tk.X, pady=(0, 6))

    preset_var = tk.StringVar(value=settings.get("preset", "auto"))
    ttk.Radiobutton(cfg_fr, text="Авто", variable=preset_var, value="auto").grid(
        row=0, column=0, sticky=tk.W, padx=4, pady=2
    )
    ttk.Radiobutton(cfg_fr, text="Только текст PDF", variable=preset_var, value="native").grid(
        row=0, column=1, sticky=tk.W, padx=4, pady=2
    )
    ttk.Radiobutton(cfg_fr, text="Всегда OCR", variable=preset_var, value="ocr_always").grid(
        row=0, column=2, sticky=tk.W, padx=4, pady=2
    )

    ttk.Label(cfg_fr, text="Порог симв.").grid(row=1, column=0, sticky=tk.W, padx=4)
    thresh_var = tk.StringVar(value=str(settings.get("min_chars", 40)))
    ttk.Entry(cfg_fr, textvariable=thresh_var, width=6).grid(row=1, column=1, sticky=tk.W)

    ttk.Label(cfg_fr, text="DPI").grid(row=1, column=2, sticky=tk.W, padx=(8, 0))
    dpi_var = tk.StringVar(value=str(settings.get("dpi", 200)))
    ttk.Entry(cfg_fr, textvariable=dpi_var, width=6).grid(row=1, column=3, sticky=tk.W)

    ttk.Label(cfg_fr, text="Языки OCR").grid(row=2, column=0, sticky=tk.W, padx=4)
    lang_var = tk.StringVar(value=settings.get("lang", "rus+eng"))
    ttk.Entry(cfg_fr, textvariable=lang_var, width=24).grid(row=2, column=1, columnspan=3, sticky=tk.W)

    pwd_var = tk.StringVar(value=settings.get("password", ""))
    ttk.Label(cfg_fr, text="Пароль PDF (если один на все)").grid(row=3, column=0, sticky=tk.W, padx=4)
    ttk.Entry(cfg_fr, textvariable=pwd_var, width=28, show="*").grid(row=3, column=1, columnspan=3, sticky=tk.W)

    adv_var = tk.BooleanVar(value=settings.get("advanced", False))
    tables_var = tk.BooleanVar(value=settings.get("extract_tables", False))
    parallel_var = tk.BooleanVar(value=settings.get("parallel", True))
    stop_err_var = tk.BooleanVar(value=settings.get("stop_on_error", False))
    pre_gray = tk.BooleanVar(value=settings.get("pre_gray", False))
    pre_bin = tk.BooleanVar(value=settings.get("pre_bin", False))

    ttk.Checkbutton(cfg_fr, text="Подробно", variable=adv_var).grid(row=4, column=0, sticky=tk.W, padx=4)
    ttk.Checkbutton(cfg_fr, text="Таблицы (pdfplumber)", variable=tables_var).grid(row=4, column=1, sticky=tk.W)
    ttk.Checkbutton(cfg_fr, text="Параллельно страницы", variable=parallel_var).grid(row=4, column=2, sticky=tk.W)

    adv_inner = ttk.Frame(cfg_fr)
    adv_inner.grid(row=5, column=0, columnspan=5, sticky=tk.W)

    ttk.Label(adv_inner, text="PSM").grid(row=0, column=0, padx=4)
    psm_var = tk.StringVar(value=str(settings.get("psm", 3)))
    ttk.Entry(adv_inner, textvariable=psm_var, width=5).grid(row=0, column=1)

    ttk.Label(adv_inner, text="OEM").grid(row=0, column=2, padx=(8, 0))
    oem_var = tk.StringVar(value=str(settings.get("oem", 3)))
    ttk.Entry(adv_inner, textvariable=oem_var, width=5).grid(row=0, column=3)

    ttk.Label(adv_inner, text="Воркеры").grid(row=0, column=4, padx=(8, 0))
    workers_var = tk.StringVar(value=str(settings.get("workers", 4)))
    ttk.Entry(adv_inner, textvariable=workers_var, width=5).grid(row=0, column=5)

    ttk.Label(adv_inner, text="Макс. стр.").grid(row=1, column=0, padx=4)
    maxp_var = tk.StringVar(value=str(settings.get("max_pages", "") or ""))
    ttk.Entry(adv_inner, textvariable=maxp_var, width=8).grid(row=1, column=1)

    ttk.Checkbutton(adv_inner, text="Стоп при ошибке", variable=stop_err_var).grid(row=1, column=2, columnspan=2)
    ttk.Checkbutton(adv_inner, text="Серый перед OCR", variable=pre_gray).grid(row=1, column=4, columnspan=2)
    ttk.Checkbutton(adv_inner, text="Бинаризация", variable=pre_bin).grid(row=2, column=0, columnspan=2)

    def toggle_adv(*_a: Any) -> None:
        if adv_var.get():
            adv_inner.grid()
        else:
            adv_inner.grid_remove()

    adv_var.trace_add("write", lambda *_: toggle_adv())
    toggle_adv()

    btn_fr = tk.Frame(root, bg=_T.CARD)
    btn_fr.pack(fill=tk.X, pady=(0, 6))

    prog = ttk.Progressbar(btn_fr, style="Hub.Horizontal.TProgressbar", mode="determinate", length=420)
    prog.pack(side=tk.LEFT, padx=(0, 8))

    stat_lbl = ttk.Label(btn_fr, text="")
    stat_lbl.pack(side=tk.LEFT)

    def build_options() -> pdf_mod.ProcessOptions:
        o = pdf_mod.ProcessOptions()
        o.preset = preset_var.get() or "auto"
        try:
            o.min_chars_native = max(0, int(thresh_var.get()))
        except ValueError:
            o.min_chars_native = 40
        try:
            o.dpi_render = max(72, min(400, int(dpi_var.get())))
        except ValueError:
            o.dpi_render = 200
        o.lang = lang_var.get().strip() or "rus+eng"
        try:
            o.psm = int(psm_var.get())
        except ValueError:
            o.psm = 3
        try:
            o.oem = int(oem_var.get())
        except ValueError:
            o.oem = 3
        mp = maxp_var.get().strip()
        if mp:
            try:
                o.max_pages = max(1, int(mp))
            except ValueError:
                o.max_pages = None
        else:
            o.max_pages = None
        o.stop_on_error = bool(stop_err_var.get())
        try:
            o.max_workers = max(1, min(16, int(workers_var.get())))
        except ValueError:
            o.max_workers = 4
        o.preprocess_grayscale = bool(pre_gray.get())
        o.preprocess_binarize = bool(pre_bin.get())
        return o

    def persist_ui_settings() -> None:
        settings["preset"] = preset_var.get()
        settings["min_chars"] = thresh_var.get()
        settings["dpi"] = dpi_var.get()
        settings["lang"] = lang_var.get()
        settings["password"] = pwd_var.get()
        settings["advanced"] = adv_var.get()
        settings["extract_tables"] = tables_var.get()
        settings["parallel"] = parallel_var.get()
        settings["stop_on_error"] = stop_err_var.get()
        settings["psm"] = psm_var.get()
        settings["oem"] = oem_var.get()
        settings["workers"] = workers_var.get()
        settings["max_pages"] = maxp_var.get()
        settings["pre_gray"] = pre_gray.get()
        settings["pre_bin"] = pre_bin.get()
        _save_settings(settings)

    def run_done(batch: pdf_mod.BatchOutcome | None, err: str | None) -> None:
        def ui() -> None:
            batch_holder[0] = batch
            clear_tab_err()
            prog["value"] = 0
            btn_run.state(["!disabled"])
            btn_cancel.state(["disabled"])
            if err:
                set_err(True)
                messagebox.showerror("Обработка PDF", err)
                stat_lbl.config(text="Ошибка")
                return
            if batch:
                n_files = len(batch.files)
                n_pages = sum(len(f.pages) for f in batch.files)
                n_ocr = sum(
                    1
                    for f in batch.files
                    for p in f.pages
                    if p.source in ("ocr", "mixed")
                )
                stat_lbl.config(text=f"Готово: {n_files} файл(ов), {n_pages} стр., OCR на {n_ocr} стр.")
                fill_preview(batch)
        try:
            родитель.after(0, ui)
        except Exception:
            pass

    def fill_preview(batch: pdf_mod.BatchOutcome) -> None:
        tv.delete(*tv.get_children())
        text_w.config(state=tk.NORMAL)
        text_w.delete("1.0", tk.END)
        buf: list[str] = []
        for fo in batch.files:
            ocr_n = sum(1 for p in fo.pages if p.source in ("ocr", "mixed"))
            tv.insert(
                "",
                tk.END,
                text=os.path.basename(fo.path),
                values=(str(fo.page_count), str(len(fo.pages)), str(ocr_n), fo.file_error or ""),
            )
            for pg in fo.pages:
                head = f"\n--- {os.path.basename(fo.path)} · стр. {pg.page_index + 1} ({pg.source}) ---\n"
                buf.append(head)
                buf.append(pg.text or "")
        text_w.insert("1.0", "".join(buf))
        text_w.config(state=tk.DISABLED)

    t0_run = [0.0]

    def do_run() -> None:
        persist_ui_settings()
        if not files_list:
            messagebox.showinfo("Обработка PDF", "Добавьте PDF файлы.")
            return
        cancel_ev.clear()
        batch_holder[0] = None
        clear_tab_err()
        btn_run.state(["disabled"])
        btn_cancel.state(["!disabled"])
        prog["value"] = 0
        stat_lbl.config(text="Обработка…")
        t0_run[0] = time.time()
        o = build_options()
        pw = pwd_var.get().strip() or None
        passwords = {os.path.abspath(p): pw for p in files_list} if pw else {}

        def on_file_start(ap: str, i: int, tot: int) -> None:
            def u() -> None:
                stat_lbl.config(text=f"Файл {i}/{tot}: {os.path.basename(ap)}")
            try:
                родитель.after(0, u)
            except Exception:
                pass

        def work() -> None:
            try:
                # подсчёт прогресса: пер-считаем страницы для прогресс бара
                prog_max = 0
                if pdf_mod.FITZ_AVAILABLE:
                    for fp in files_list:
                        try:
                            d, e = pdf_mod.open_pdf_authenticated(fp, pw)
                            if d:
                                n = d.page_count
                                if o.max_pages:
                                    n = min(n, o.max_pages)
                                prog_max += n
                                d.close()
                        except Exception:
                            prog_max += 1
                if prog_max < 1:
                    prog_max = len(files_list)

                done_pages = [0]

                def page_cb(_ap: str, cur: int, tot: int) -> None:
                    done_pages[0] += 1

                    def u2() -> None:
                        prog["maximum"] = max(1, prog_max)
                        prog["value"] = min(prog_max, done_pages[0])
                        elapsed = time.time() - t0_run[0]
                        rate = done_pages[0] / elapsed if elapsed > 0.01 else 0
                        rest = prog_max - done_pages[0]
                        eta = rest / rate if rate > 0 else 0
                        stat_lbl.config(text=f"Страница {done_pages[0]}/{prog_max} · ~{int(eta)} с")

                    try:
                        родитель.after(0, u2)
                    except Exception:
                        pass

                b = pdf_mod.process_paths(
                    list(files_list),
                    o,
                    passwords,
                    cancel_ev,
                    extract_tables=bool(tables_var.get()),
                    parallel_pages=bool(parallel_var.get()),
                    on_file_start=on_file_start,
                    on_page_progress=page_cb,
                )
                run_done(b, None)
            except Exception as e:
                run_done(None, str(e))

        th = threading.Thread(target=work, daemon=True)
        worker[0] = th
        th.start()

    def cancel_run() -> None:
        cancel_ev.set()
        stat_lbl.config(text="Отмена…")

    btn_run = ttk.Button(btn_fr, text="Обработать", style="Accent.TButton", command=do_run)
    btn_run.pack(side=tk.LEFT, padx=(0, 6))
    btn_cancel = ttk.Button(btn_fr, text="Отмена", command=cancel_run, state=tk.DISABLED)
    btn_cancel.pack(side=tk.LEFT, padx=(0, 12))

    export_fr = tk.Frame(root, bg=_T.CARD)
    export_fr.pack(fill=tk.X, pady=(4, 0))

    def export_wrap(kind: str) -> None:
        b = batch_holder[0]
        if not b:
            messagebox.showinfo("Обработка PDF", "Сначала выполните «Обработать».")
            return
        init = settings.get("last_export_dir") or settings.get("last_dir") or os.path.expanduser("~")
        ts = time.strftime("%Y-%m-%d_%H-%M")
        if kind == "xlsx":
            p = filedialog.asksaveasfilename(
                title="Excel",
                defaultextension=".xlsx",
                initialfile=f"ОбработкаPDF_{ts}.xlsx",
                initialdir=init if os.path.isdir(init) else None,
                filetypes=[("Excel", "*.xlsx")],
            )
            if p:
                settings["last_export_dir"] = os.path.dirname(p)
                _save_settings(settings)
                ok, msg = export_excel(p, b)
                if ok:
                    messagebox.showinfo("Обработка PDF", "Файл Excel сохранён.")
                else:
                    messagebox.showerror("Обработка PDF", msg)
        elif kind == "docx":
            p = filedialog.asksaveasfilename(
                title="Word",
                defaultextension=".docx",
                initialfile=f"ОбработкаPDF_{ts}.docx",
                initialdir=init if os.path.isdir(init) else None,
                filetypes=[("Word", "*.docx")],
            )
            if p:
                settings["last_export_dir"] = os.path.dirname(p)
                _save_settings(settings)
                short = messagebox.askyesno("Обработка PDF", "Только краткие превью по страницам?")
                ok, msg = export_word(p, b, short)
                if ok:
                    messagebox.showinfo("Обработка PDF", "Файл Word сохранён.")
                else:
                    messagebox.showerror("Обработка PDF", msg)
        elif kind == "json":
            p = filedialog.asksaveasfilename(
                title="JSON",
                defaultextension=".json",
                initialfile=f"ОбработкаPDF_{ts}.json",
                initialdir=init if os.path.isdir(init) else None,
                filetypes=[("JSON", "*.json")],
            )
            if p:
                settings["last_export_dir"] = os.path.dirname(p)
                _save_settings(settings)
                ok, msg = export_json(p, b)
                if ok:
                    messagebox.showinfo("Обработка PDF", "JSON сохранён.")
                else:
                    messagebox.showerror("Обработка PDF", msg)
        elif kind == "csv":
            p = filedialog.asksaveasfilename(
                title="CSV",
                defaultextension=".csv",
                initialfile=f"ОбработкаPDF_{ts}.csv",
                initialdir=init if os.path.isdir(init) else None,
                filetypes=[("CSV", "*.csv")],
            )
            if p:
                settings["last_export_dir"] = os.path.dirname(p)
                _save_settings(settings)
                ok, msg = export_csv(p, b)
                if ok:
                    messagebox.showinfo("Обработка PDF", "CSV сохранён.")
                else:
                    messagebox.showerror("Обработка PDF", msg)
        elif kind == "search_pdf":
            p = filedialog.asksaveasfilename(
                title="Searchable PDF",
                defaultextension=".pdf",
                initialfile=f"Searchable_{ts}.pdf",
                initialdir=init if os.path.isdir(init) else None,
                filetypes=[("PDF", "*.pdf")],
            )
            if not p:
                return
            settings["last_export_dir"] = os.path.dirname(p)
            _save_settings(settings)
            if len(files_list) != 1:
                messagebox.showerror("Обработка PDF", "Выберите ровно один исходный PDF в списке (для объединённого экспорта за один раз).")
                return
            opts = build_options()
            ok, msg = pdf_mod.build_searchable_pdf(
                os.path.abspath(files_list[0]),
                p,
                opts,
                pwd_var.get().strip() or None,
            )
            if ok:
                messagebox.showinfo("Обработка PDF", "PDF сохранён (экспериментальный режим поиска).")
            else:
                messagebox.showerror("Обработка PDF", msg)

    ttk.Button(export_fr, text="Excel…", command=lambda: export_wrap("xlsx")).pack(side=tk.LEFT, padx=(0, 4))
    ttk.Button(export_fr, text="Word…", command=lambda: export_wrap("docx")).pack(side=tk.LEFT, padx=(0, 4))
    ttk.Button(export_fr, text="JSON…", command=lambda: export_wrap("json")).pack(side=tk.LEFT, padx=(0, 4))
    ttk.Button(export_fr, text="CSV…", command=lambda: export_wrap("csv")).pack(side=tk.LEFT, padx=(0, 4))
    ttk.Button(export_fr, text="Searchable PDF…", command=lambda: export_wrap("search_pdf")).pack(side=tk.LEFT, padx=(0, 4))

    hint = tk.Text(
        root,
        height=3,
        font=_T.FONT_SM,
        fg=_T.TEXT_DIM,
        bg=_T.CARD_ALT,
        relief=tk.FLAT,
        wrap=tk.WORD,
    )
    hint.pack(fill=tk.X, pady=(8, 0))
    deps = []
    if not pdf_mod.FITZ_AVAILABLE:
        deps.append("pymupdf")
    if not pdf_mod.TESSERACT_AVAILABLE:
        deps.append("pytesseract + Tesseract OCR (см. УСТАНОВКА.txt)")
    hint.insert(
        tk.END,
        "Совет: для OCR установите Tesseract для Windows и языки rus+eng. "
        + (" Отсутствует: " + ", ".join(deps) if deps else "Зависимости в порядке."),
    )
    hint.config(state=tk.DISABLED)

    nb = ttk.Notebook(root)
    nb.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

    tab_sum = ttk.Frame(nb)
    nb.add(tab_sum, text="Сводка")
    tv = ttk.Treeview(
        tab_sum,
        columns=("pages", "proc", "ocr", "err"),
        show="tree headings",
        height=8,
    )
    tv.heading("#0", text="Файл")
    tv.heading("pages", text="Стр. в PDF")
    tv.heading("proc", text="Обработано стр.")
    tv.heading("ocr", text="С OCR")
    tv.heading("err", text="Ошибка файла")
    tv.pack(fill=tk.BOTH, expand=True)

    tab_txt = ttk.Frame(nb)
    nb.add(tab_txt, text="Текст")
    sx = ttk.Scrollbar(tab_txt)
    sx.pack(side=tk.RIGHT, fill=tk.Y)
    text_w = tk.Text(
        tab_txt,
        wrap=tk.WORD,
        font=_T.FONT_MONO,
        yscrollcommand=sx.set,
        bg=_T.CARD_ALT,
        fg=_T.TEXT,
        relief=tk.FLAT,
        state=tk.DISABLED,
    )
    text_w.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    sx.config(command=text_w.yview)

    refresh_file_list()
