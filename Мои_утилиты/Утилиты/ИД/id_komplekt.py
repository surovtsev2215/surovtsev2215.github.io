# -*- coding: utf-8 -*-
"""Формирование комплекта ИД по образцу: реестр (форма 1.2), титул, опционально копирование всей папки образца."""

from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from openpyxl import load_workbook


def _cfg_path(util_dir: Path) -> Path:
    return util_dir / "Данные" / "комплект_id.json"


def загрузить_конфиг(util_dir: Path) -> dict[str, Any]:
    p = _cfg_path(util_dir)
    if not p.is_file():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def сохранить_конфиг(util_dir: Path, data: dict[str, Any]) -> None:
    p = _cfg_path(util_dir)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(p)


def _найти_строку_примечания(ws, default: int = 26) -> int:
    for r in range(20, 55):
        v = ws.cell(row=r, column=2).value
        if v and isinstance(v, str) and "Исполнительная документация составлена" in v:
            return r
    return default


def заполнить_реестр_xlsx(шаблон_xlsx: Path, выход_xlsx: Path, cfg: dict[str, Any]) -> None:
    wb = load_workbook(шаблон_xlsx)
    ws = wb["1439.1"] if "1439.1" in wb.sheetnames else wb.active

    шифр = str(cfg.get("шифр") or "").strip()
    ws["A4"] = cfg.get("заказчик_блок") or ""
    ws["F4"] = cfg.get("строительство_блок") or ""
    ws["F13"] = f"Проект: {шифр}" if шифр else ""
    ws["B16"] = f"Шифр проекта {шифр}" if шифр else ""
    ks = str(cfg.get("основание_кс") or "").strip()
    ws["B17"] = f"Основание: {ks}" if ks else ""

    entries = cfg.get("позиции_реестра") or []
    org_def = str(cfg.get("организация_исполнитель") or "").strip()
    note_row = _найти_строку_примечания(ws)
    old_slots = note_row - 22
    need = len(entries)
    if need > old_slots:
        ws.insert_rows(note_row, need - old_slots)
        note_row += need - old_slots

    for r in range(22, note_row):
        for col in (1, 2, 3, 13, 14):
            ws.cell(row=r, column=col, value=None)

    for i, row in enumerate(entries):
        r = 22 + i
        ws.cell(row=r, column=1, value=i + 1)
        ws.cell(row=r, column=2, value=str(row.get("наименование") or ""))
        ws.cell(row=r, column=3, value=str(row.get("обозначение") or ""))
        og = str(row.get("организация") or "").strip() or org_def
        ws.cell(row=r, column=13, value=og)
        л = row.get("листов")
        ws.cell(row=r, column=14, value=л if л not in ("", None) else None)

    note_txt = cfg.get("примечание_после_реестра")
    if note_txt:
        ws.cell(row=note_row, column=2, value=str(note_txt))

    выход_xlsx.parent.mkdir(parents=True, exist_ok=True)
    wb.save(выход_xlsx)
    wb.close()


def _заменить_в_тексте(text: str, пары: list[list[str]]) -> str:
    for pair in пары:
        if len(pair) >= 2:
            old, new = str(pair[0]), str(pair[1])
            if old:
                text = text.replace(old, new)
    return text


def заполнить_титул_docx(шаблон_docx: Path, выход_docx: Path, cfg: dict[str, Any]) -> None:
    from docx import Document

    пары = cfg.get("замены_в_титуле") or []
    if not isinstance(пары, list):
        пары = []

    doc = Document(str(шаблон_docx))

    def обработать_абзацы(paragraphs) -> None:
        for p in paragraphs:
            for run in p.runs:
                run.text = _заменить_в_тексте(run.text, пары)

    обработать_абзацы(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                обработать_абзацы(cell.paragraphs)

    выход_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(выход_docx))


def сформировать_комплект(
    util_dir: Path,
    папка_выхода: Path,
    cfg: dict[str, Any] | None = None,
    manifest: dict[str, Any] | None = None,
) -> tuple[bool, str]:
    """
    Если ``копировать_все_файлы_из_образца`` и задана ``папка_образца`` — копируется всё содержимое образца,
    затем перезаписываются реестр и титул по шаблонам из ``util_dir/Шаблоны``.
    Иначе создаются только титул и реестр в выходной папке.

    Служебные ключи конфига (начинаются с «_») не передаются в Excel/Word.
    Если задан ``manifest``, сохраняется файл ``manifest_id.json`` в выходной папке.
    """
    cfg = cfg if cfg is not None else загрузить_конфиг(util_dir)
    cfg_fill = {k: v for k, v in cfg.items() if not str(k).startswith("_")}
    шаблоны = util_dir / "Шаблоны"
    tpl_xlsx = next(шаблоны.glob("*форма12*.xlsx"), None)
    tpl_docx = шаблоны / "титул.docx"
    if tpl_xlsx is None or not tpl_xlsx.is_file():
        return False, "Нет шаблона реестра в папке Утилиты/ИД/Шаблоны (ожидается *форма12*.xlsx)."
    if not tpl_docx.is_file():
        tpl_docx = next(шаблоны.glob("*.docx"), None)
    if tpl_docx is None or not Path(tpl_docx).is_file():
        return False, "Нет шаблона титула (.docx) в Шаблоны."

    имя_реестра = str(cfg_fill.get("имя_реестра") or "01. Реестр исполнительной документации.xlsx")
    имя_титула = str(cfg_fill.get("имя_титула") or "0. Титульный лист.docx")

    образец = Path(str(cfg_fill.get("папка_образца") or "").strip())
    копировать = bool(cfg_fill.get("копировать_все_файлы_из_образца")) and образец.is_dir()

    папка_выхода.mkdir(parents=True, exist_ok=True)

    if копировать:
        for item in образец.iterdir():
            dest = папка_выхода / item.name
            if item.is_dir():
                if dest.exists():
                    shutil.rmtree(dest)
                shutil.copytree(item, dest)
            else:
                shutil.copy2(item, dest)
    else:
        if not any(папка_выхода.iterdir()):
            pass

    out_xlsx = папка_выхода / имя_реестра
    out_docx = папка_выхода / имя_титула

    try:
        заполнить_реестр_xlsx(tpl_xlsx, out_xlsx, cfg_fill)
    except Exception as e:
        return False, f"Ошибка заполнения реестра: {e}"

    try:
        заполнить_титул_docx(Path(tpl_docx), out_docx, cfg_fill)
    except Exception as e:
        return False, f"Ошибка заполнения титула: {e}"

    msg = "Комплект сформирован: реестр и титул записаны в выбранную папку."
    if копировать:
        msg += " Скопированы все файлы из папки образца; реестр и титул заменены заполненными."
    if manifest:
        try:
            mf = папка_выхода / "manifest_id.json"
            mf.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
            msg += " Добавлен manifest_id.json (контрольные суммы и проверки)."
        except Exception as e:
            return False, f"Комплект записан, но manifest не сохранён: {e}"
    return True, msg
